import json
import os.path
import time
import uuid
import re
import PyPDF2
import pandas as pd
import docx
from sentence_transformers import SentenceTransformer
import tiktoken
import asyncio
import pdfplumber
import chromadb
from chromadb.config import DEFAULT_TENANT, DEFAULT_DATABASE, Settings


persistent_storage_path = "./chromadb_storage"  # Change this path if needed
# Create the directory if it does not exist
if not os.path.exists(persistent_storage_path):
    os.makedirs(persistent_storage_path)
    

# Initialize ChromaDB client with a persistent storage path
chroma_client = chromadb.PersistentClient(
    path=persistent_storage_path,
    settings=Settings(),
    tenant=DEFAULT_TENANT,
    database=DEFAULT_DATABASE,
)

collection = chroma_client.get_or_create_collection("document_embeddings")

model = SentenceTransformer(model_name_or_path='all-MiniLM-L6-v2', device='cpu')

def count_tokens(text: str, encoding_name: str = 'cl100k_base'):
    if not isinstance(text, str) or text is None:
        raise ValueError("Input text must be a non-empty string.")
    encoding = tiktoken.get_encoding(encoding_name)
    tokens = encoding.encode(text)
    num_tokens = len(tokens)
    return num_tokens

def simple_token_count(text: str):
    tokens = re.findall(r"\w+|[^\w\s]", text, re.UNICODE)
    return len(tokens)

def clean_text(text):
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    special_unicode_pattern = r'\\u[\dA-Fa-f]{4}|\u200b'
    text = text.replace('\n', ' ')
    text = re.sub(special_unicode_pattern, '', text)
    text = text.replace('\\', '')
    urls = re.findall(url_pattern, text)
    for idx, url in enumerate(urls):
        text = text.replace(url, f"URLPLACEHOLDER{idx}", 1)
    emails = re.findall(email_pattern, text)
    for idx, email in enumerate(emails):
        text = text.replace(email, f"EMAILPLACEHOLDER{idx}", 1)
    text = re.sub(r'\s+', ' ', text)
    for idx, url in enumerate(urls):
        text = text.replace(f"URLPLACEHOLDER{idx}", url, 1)
    for idx, email in enumerate(emails):
        text = text.replace(f"EMAILPLACEHOLDER{idx}", email, 1)
    return text

def extract_and_embed_links(page):
    text_with_links = page.extract_text() or ""
    annotations = page.annots

    if annotations:
        for annot in annotations:
            if annot is not None and 'uri' in annot and annot['uri']:
                uri = annot['uri']
                # Skip email addresses
                if uri.startswith("mailto:"):
                    continue

                try:
                    words = page.extract_words()
                    # Print the structure of the word dictionary for debugging
                    if words:
                        print(f"Word structure: {words[0]}")

                    # Extracting text based on link location
                    linked_words = [word for word in words if
                                    annot['doctop'] <= word['bottom'] and
                                    annot['bottom'] >= word['top'] and
                                    annot['x0'] <= word['x1'] and
                                    annot['x1'] >= word['x0']]
                    linked_text = " ".join(word['text'] for word in linked_words)
                    link_str = f"{linked_text} ({uri})"

                    if linked_text:
                        text_with_links = text_with_links.replace(linked_text, link_str)
                except Exception as e:
                    print(f"Error processing link for text: {str(e)}")

    return text_with_links

def chunk_document(text, metadata=None, max_tokens=7500):
    current_chunk = ""
    meta = ""
    current_tokens = 0
    num_lines = 0

    if metadata:
        meta += f"Title: {metadata['title']}\n"
        meta += f"Document File Name: {metadata['filename']}\n"
        meta += "\n\n"
        current_chunk += meta
        current_tokens += count_tokens(current_chunk)
        num_lines += 1

    for line in text.split('\n'):
        try:
            line_tokens = count_tokens(line)
        except ValueError:
            line_tokens = simple_token_count(line)
        if current_tokens + line_tokens > max_tokens:
            yield current_chunk + line + '\n'
            current_chunk = meta + line + '\n'
            current_tokens = line_tokens
            num_lines = 1
        else:
            current_chunk += line + '\n'
            current_tokens += line_tokens
            num_lines += 1

        if num_lines >= max_tokens:
            yield current_chunk + line + '\n'
            current_chunk = ""
            current_tokens = 0
            num_lines = 0

    if current_chunk:
        yield current_chunk

async def embeddings_and_storage(text, metadata, max_tokens):
    doc_uuid = str(uuid.uuid4())

    # Iterate over chunks of the document
    for i, chunk in enumerate(chunk_document(text, metadata, max_tokens)):
        try:
            chunk_token = count_tokens(chunk)
        except ValueError:
            chunk_token = simple_token_count(chunk)
        print(f"Chunk: {i} with Token: {chunk_token}", flush=True)
        await asyncio.sleep(1)

        # Generate embeddings for the chunk
        chunk_embed = model.encode([chunk])[0]  # Encoding a single chunk
        chunk_uuid = str(uuid.uuid4())  # Unique ID for the chunk

        # Store the chunk in ChromaDB with metadata and unique ID
        collection.add(
            ids=[chunk_uuid],  # Unique ID for the embedding entry
            embeddings=[chunk_embed.tolist()],  # List of embeddings
            metadatas=[{
                'doc_uuid': doc_uuid,
                'chunk_uuid': chunk_uuid,
                'title': metadata.get('title', metadata.get('filename')),
                'doc_filename': metadata['filename'],
                'chunk': chunk
            }],
            documents=[chunk]
        )

        print(f"Chunk {i} stored in ChromaDB.", flush=True)

    return "ChromaDB storage complete."

async def read_pdf(filename, max_tokens):
    """Read a PDF file, extract text with hyperlinks inserted next to the relevant text, and store the results."""
    full_text_with_links = ""
    doc_name = filename.split('/')[-1]
    with pdfplumber.open(filename) as pdf:
        for page in pdf.pages:
            full_text_with_links += extract_and_embed_links(page) + "\n"

    # Process for storing in CSV
    metadata = {
        'title': pdf.metadata.get('Title', doc_name),
        'filename': doc_name,
        'total_pages': len(pdf.pages)
    }
    text = full_text_with_links.replace('\n', ' ')
    await embeddings_and_storage(text, metadata, max_tokens)

    return full_text_with_links

async def read_json(file_path):
    file_name = file_path.split('/')[-1]
    with open(file_path) as json_file:
        json_data = json.load(json_file)
        for entry in json_data:
            title = entry.get('title', 'Unknown Title')
            doc_data = entry.get('doc_data', "")
            metadata = {
                'filename': file_name,
                'title': title
            }
            text = doc_data.replace('\n', ' ')
            await embeddings_and_storage(text, metadata, 7500)

async def read_excel_cell_data(file_path):
    df = pd.read_excel(file_path)
    for index, row in df.iterrows():
        await asyncio.sleep(2)
        print(f"Writing doc title: {row['title']}")
        doc_data = row['doc_data']
        doc_chunks = chunk_document(doc_data)
        doc_uuid = str(uuid.uuid4())
        with open('output.csv', mode='a', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['doc_uuid', 'chunk_uuid', 'title', 'doc_filename', 'chunk', 'embedding']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            for i, chunk in enumerate(doc_chunks):
                chunk_token = count_tokens(chunk)
                print(f"Chunk: {i} with Token: {chunk_token}")
                await asyncio.sleep(1)
                chunk_embed = model.encode([chunk])[0]
                chunk_uuid = str(uuid.uuid4())

                row = {
                    'doc_uuid': doc_uuid,
                    'chunk_uuid': chunk_uuid,
                    'title': row["title"],
                    'doc_filename': row["doc_link"],
                    'chunk': chunk,
                    'embedding': json.dumps(chunk_embed.tolist())  # Convert embedding to list and then to JSON string
                }
                writer.writerow(row)
        print("Embeddings and Document Successfully Stored in CSV")
    return 'success'

async def read_docx(filename, max_tokens):
    """Read a docx file, collect metadata and yield chunks of text."""
    doc = docx.Document(filename)
    doc_name = filename.split('/')[-1]
    title = doc.paragraphs[0].text.strip()
    metadata = {
        'filename': doc_name,
        'title': title,
        'total_pages': len(doc.paragraphs),
    }
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    text = text.replace('\n', ' ')
    await embeddings_and_storage(text, metadata, max_tokens)


async def process_file_by_type(file_path, max_tokens=7500):
    """
    This function automatically detects the file type based on the file extension
    and processes the file accordingly (PDF, DOCX, Excel, or JSON).
    
    Args:
        file_path (str): Path to the file to be processed.
        max_tokens (int): Maximum tokens allowed per chunk of text.
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.pdf':
        print(f"Processing PDF file: {file_path}", flush=True)
        await read_pdf(file_path, max_tokens)
    elif file_extension == '.docx':
        print(f"Processing DOCX file: {file_path}", flush=True)
        await read_docx(file_path, max_tokens)
    elif file_extension == '.xlsx':
        print(f"Processing Excel file: {file_path}", flush=True)
        await read_excel_cell_data(file_path)
    elif file_extension == '.json':
        print(f"Processing JSON file: {file_path}", flush=True)
        await read_json(file_path)
    else:
        print(f"Unsupported file type: {file_extension}. Please provide a PDF, DOCX, Excel, or JSON file.", flush=True)


async def list_all_data_in_chromadb():
    """List all data stored in the ChromaDB collection."""
    # Retrieve all documents in the collection
    all_documents = collection.get()

    if not all_documents:
        print("No data found in ChromaDB.")
        return

    print("Listing all data in ChromaDB:\n")
    for i, document in enumerate(all_documents['documents']):
        metadata = all_documents['metadatas'][i]
        print(f"Document {i+1}:")
        print(f"Title: {metadata.get('title', 'No title')}")
        print(f"Filename: {metadata.get('doc_filename', 'No filename')}")
        print(f"Chunk UUID: {metadata.get('chunk_uuid', 'No chunk UUID')}")
        print(f"Doc UUID: {metadata.get('doc_uuid', 'No doc UUID')}")
        print(f"Chunk Text: {document[:200]}...")  # Display the first 200 characters for preview
        print("\n---\n")
        
        
        
        
        
async def main():
    file_path = "./test.pdf"
    await process_file_by_type(file_path)
    await list_all_data_in_chromadb()

if __name__ == "__main__":
    asyncio.run(main())